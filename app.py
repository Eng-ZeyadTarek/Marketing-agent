import streamlit as st
from docx import Document as DocxDocument  # Import from docx and rename to avoid conflict
import io
import os
import openai

from langchain_community.utilities import SerpAPIWrapper
from langchain_openai import ChatOpenAI
from langchain.chains import RetrievalQA
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
import openai
from langchain_core.prompts import PromptTemplate
from langchain.schema import Document
from langchain.agents.tools import Tool
from langchain.agents import AgentExecutor, create_openai_tools_agent
from langchain import hub
import streamlit as st

# Define the CSS to inject
def local_css(file_name):
    with open(file_name) as f:
        st.markdown(f'<style>{f.read()}</style>', unsafe_allow_html=True)

# Inject CSS with the specified file
#local_css(r"C:\Users\zeyad\Documents\Marketing_bot_project\style.css")

def read_word_document(file_buffer):
    doc = DocxDocument(io.BytesIO(file_buffer.getvalue()))
    data = []  # Store question and answer pairs
    current_question = None
    answer_lines = []
    
    paragraphs = [p.text for p in doc.paragraphs]

    questions = [
        "What is the Dilemma? What is the opportunity out there? What is the gap that be filled?",
        "How will your brand solve the Dilemma? What does the brand promise the consumer?",
        "What products/service do you offer to solve this?",
        "Why should people believe it? What makes you credible?",
        "What are the trends in the category that inspired you?",
        "Who inspires you in this or similar categories? Please list their website domains",
        "Who are your main competitors (list 3-5)? What is their strength, weakness and competitive edge? Please list their website domain",
        "What are the demographics/psychographics of your audience?",
        "Who are the audience that define/personify your brand the most? Who will readily buy the brand (early adopters)?",
        "Who are the audience that would never personify your brand?",
        "Who would represent/endorse your brand? Describe why?",
        "What are their pain/desires? What do they need out of the category?",
        "Who do you want to convince into buying your brand?",
        "How can we best reach them (main channels)?",
        "What is your brand purpose & values?",
        "What do you want people to say about your brand? What is the aspired brand personality? Which words do you want to own?",
        "Which traits will never define your brand?",
        "Is there any word, image, phrase or story that reflects what the brand stands for?",
        "Is there anything else you’d like to add?",
        "What are you expecting out of this branding exercise?",
        "List all your products/services/special features and the respective grouping if relevant.",
        "List any additions or extensions to your product/service that you might introduce in the future.",
    ]
    # Add headers or other questions that also indicate a new section here
    headers = [
        "The Brand Proposition",
        "The Category",
        "The Competition",
        # Add all other headers here...
        "The Brand Purpose & Perception",
    ]

    def is_question(paragraph):
        return any(paragraph.strip() == q for q in questions)

    def is_header(paragraph):
        return any(paragraph.strip() == h for h in headers)

    for i, paragraph in enumerate(paragraphs):
        paragraph_text = paragraph.strip()
        # Check for illustration in the next paragraph

        
        next_paragraph_text = paragraphs[i + 1].strip() if i + 1 < len(paragraphs) else ""

        if is_question(paragraph_text) or (current_question and is_question(paragraph_text + " " + next_paragraph_text)):
            if current_question:
                # Save the current Q&A, ensuring not to duplicate the illustration in the answer
                data.append((current_question, '\n'.join(answer_lines)))
                answer_lines = []
            current_question = paragraph_text
            if next_paragraph_text and not is_question(next_paragraph_text) and not is_header(next_paragraph_text):
                current_question += " " + next_paragraph_text  # Append illustration
                i += 1  # Skip the next paragraph as it's been included
        elif current_question and not is_header(paragraph_text):
            # Ensure not to immediately repeat the illustration in the answer
            if paragraph_text not in current_question:
                answer_lines.append(paragraph.strip())

    # Handle the last question and its answer
    if current_question and answer_lines:
        data.append((current_question, '\n'.join(answer_lines)))

    # Remove any empty Q&A caused by headers
    data = [(q, a) for q, a in data if q and a]

    return data


st.title('Marketing Strategy and Campaign Generator')

openai_key = st.text_input("Enter your OpenAI API Key", type="password")
serpapi_key = st.text_input("Enter your SerpAPI Key", type="password")
splitted_documents = []
if st.button('Confirm API Keys'):
    if openai_key and serpapi_key:
        # Store API keys in session state and mark as confirmed
        st.session_state['openai_key'] = openai_key
        st.session_state['serpapi_key'] = serpapi_key
        st.session_state['keys_confirmed'] = True
        

        st.success('API Keys confirmed. Please proceed.')
        os.environ["SERPAPI_API_KEY"] = st.session_state['serpapi_key']
        os.environ["OPENAI_API_KEY"] = st.session_state['openai_key']
        openai.api_key = st.session_state['openai_key']
        params = {
            "engine": "google",
            "gl": "eg",
            "hl": "ar",
        }
        search = SerpAPIWrapper(params=params)

    else:
        st.error("Please enter both OpenAI and SerpAPI keys to proceed.")

if st.session_state.get('keys_confirmed', False):
    uploaded_file = st.file_uploader("Choose a Word file", type=["docx"])
    BRAND_NAME = st.text_input("Enter your Brand Name")
    INDUSTRY = st.text_input("Enter your Industry")
    
    if st.button('Generate Marketing Strategy and Campaign'):
        
        params = {
                "engine": "google",
                "gl": "eg",
                "hl": "ar",
            }
        search = SerpAPIWrapper(params=params)

        if uploaded_file and BRAND_NAME and INDUSTRY:
            # Set the API keys in the environment
            
            #st.write(f"Brand Name: {brand_name}")
            #st.write(f"Industry: {industry}")

            document_content = read_word_document(uploaded_file)
            
            
            for question, answer in document_content:
                splitted_documents.append(Document(page_content=question + "\n" + answer + "\n\n"))
                #st.subheader(question)
                #st.write(answer)
            
            turbo = ChatOpenAI(model= "gpt-3.5-turbo-16k", temperature=0.0, max_tokens=4000, streaming = False)
            main_model = ChatOpenAI(model= "gpt-4", temperature=0.05, max_tokens=1000, streaming = False)
            
            embeddings = OpenAIEmbeddings(chunk_size=500)
            marketing_retriever = Chroma.from_documents(splitted_documents, embeddings, collection_name="brand")
            marketing_qa_template = """ the user is a marketing agent and he needs answers to some questions about your brand, the questions, and answers in the below context, so provide the answer to his question from the below context as it is in an accurate way, and don't make up answers or add info to the context answer from your mind.

            ===========================

            {context}

            Question: {question}"""

            MARKETING_PROMPT = PromptTemplate(
                template=marketing_qa_template, input_variables=["context", "question"]
            )
            marketing_qa = RetrievalQA.from_chain_type(
                llm=turbo,
                chain_type="stuff",
                retriever=marketing_retriever.as_retriever(search_kwargs={'k': 2}),
                chain_type_kwargs={"prompt": MARKETING_PROMPT, 'verbose': False},
                
            )
            questionaire_report = ""
            with st.spinner('Reading uploaded document...'):

            
                for i, q_a in enumerate(document_content):
                    questionaire_report += f"Q{i + 1}- " + q_a[0] +"\n"+ marketing_qa.invoke(q_a[0])['result'] +"\n\n"
            
            prompt = hub.pull("hwchase17/openai-tools-agent")
            
            tools = [Tool(name = "brand_answers", func =lambda q:str(marketing_qa({"query": q})['result']), description = "You must use it to get info about the brand that you'll create a marketing strategy for it.",
            handle_tool_error = True),
                    Tool(name = "search", description="Use it to search online to gather information about the industry and the trends", func=search.run, handle_tool_error = True)
            ]
            
            prompt.messages[0].prompt.template = f"""Upon receiving the command "start", you, as the specialized scout for '{BRAND_NAME}' in the '{INDUSTRY}', are tasked with a crucial mission. Without needing further input, you will autonomously navigate through two pivotal phases to develop a comprehensive market report. Your journey unfolds as follows:

                    - Automatically invoke the "brand_answers" tool to ask the following questions:
                    1. What products/service do you offer?
                    2. Who are your main competitors?

                    Main answer and final report: Conducting In-depth Market Analysis
                    - With the key information in hand, utilize the "search" tool to delve into:
                    1. Current market trends and analysis within the '{INDUSTRY}', identifying opportunities and threats.
                    2. Detailed information of each competitor (search for one competitor at a time) and put `company` after the company name, revealing potential advantages for '{BRAND_NAME}'.
                    3. Detailed strengths and weaknesses of each competitor (search for one competitor at a time) and put `company` after the company name.
                    4. Pricing list for competitors' products/services, aiming to pinpoint market positioning opportunities.

                    Your objective is to synthesize the information of 'search' tool into a very detailed well-structured report that encompasses market trends, competitive analysis, and pricing insights. don't put recommendations on.

                    Embark on this mission with the initiative, utilizing your tools to their fullest potential without further prompts. Your findings will forge the path for '{BRAND_NAME}' to navigate the competitive landscape and seize market opportunities with precision and insight."""

            sec_agent = create_openai_tools_agent(turbo, tools, prompt)
            sec_agent_executor = AgentExecutor(agent=sec_agent, tools=tools, verbose=True, max_iterations = 500, max_tokens = 3000)
            
            online_report = ""
            with st.spinner('Searching online...'):
                online_report = sec_agent_executor.invoke(({"input":"start"}))['output']
            questionaire_report += "\n\n" + online_report
            


            prompt.messages[0].prompt.template = f"""As a marketing bot armed with the detailed insights from {BRAND_NAME}'s survey responses and the comprehensive online report, your mission is to construct a Matter Pyramid that captures the brand's core identity within the {INDUSTRY}. Your output will be a pyramid divided into three tiers, each one echoing a key aspect of the brand's DNA: Functional Benefit, Culture and Values, and Emotional Benefit.

            1. **Functional Benefit**: Delve into the brand's offerings and the unique solutions it provides for customer challenges. From this, determine and define the foundational differentiators that establish {BRAND_NAME} as a leader in the {INDUSTRY}. Summarize this under 'Functional Benefit,' showcasing the intelligent decisions the brand makes regarding environmental stewardship, investment value, and familial care.

            2. **Culture and Values**: Reflect on {BRAND_NAME}'s internal ethos, its commitments, and the values it upholds. Under 'Culture and Values,' draft a narrative that puts the brand's consideration for its partners and stakeholders at the forefront, emphasizing a commitment to adding value and enriching every interaction.

            3. **Emotional Benefit**: Glean from the survey and report the emotional threads that bind the customers to the brand. In 'Emotional Benefit,' narrate how {BRAND_NAME} offers peace of mind, illustrating a deep understanding of customer needs and the trust they place in the brand, ensuring they are perpetually 'in safe hands.'

            Craft a response that creatively distills these three tiers into a compelling and brief narrative. This narrative should align with {BRAND_NAME}'s market position and articulate the unique story of how it stands out in the {INDUSTRY}."""

            
            main_tools  = [Tool(name = "nothing", func =search.run, description = "never use me",
            handle_tool_error = True) ]
            main_model.max_tokens = 2000
            agent_1 = create_openai_tools_agent(main_model,main_tools, prompt)
            agent_executor_1 = AgentExecutor(agent=agent_1, tools=main_tools, verbose=True, max_iterations = 500, max_tokens = 3000)
            with st.spinner('Creating matter pyramid...'):
                matter_pyramid = agent_executor_1.invoke(({"input":"start"}))['output']
                
            prompt.messages[0].prompt.template = f"""You are a brand called '{BRAND_NAME}', focusing on innovation and leadership within the '{INDUSTRY}' sector. you have been provided answers to the questionnaire and an online report for your brand, and use all the information within it to deliver a very long, very detailed, and creative marketing strategy and campaign about yourself as a cohesive narrative that not only outlines the tactical approach but also tells the story of '{BRAND_NAME}' and its journey to redefine '{BRAND_NAME}' in its industry.

            1. **Brand Insights**: Deeply analyze the input on {BRAND_NAME}'s market dilemma, solutions, product offerings, and credibility. Identify the core opportunities Ivy aims to capture and how it differentiates itself in addressing customer needs, and present it like you are talking about youself.

            2. **Competitor Analysis**: Examine {BRAND_NAME}'s main competitors, their strengths, weaknesses, and market positions. Highlight Ivy's competitive advantage based on this analysis.

            3. **Target Audience**: Define {BRAND_NAME}'s primary and secondary target audiences based on the demographics and psychographics provided. Outline strategies for engaging these audiences.

            4. **Marketing Strategy**:
                - **Brand Positioning**: Summarize {BRAND_NAME}'s market position and unique value proposition.
                - **Engagement Channels**: Identify the most effective channels for reaching {BRAND_NAME}'s target audience, considering both digital and traditional platforms.
                - **Content and Messaging**: Suggest key messages that resonate with {BRAND_NAME}'s brand promise and audience's expectations.

            5. **Campaign Concept**: Propose a campaign that embodies {BRAND_NAME}'s strategic goals, including a creative titles, objectives, key activities, and expected outcomes.

            6. **Implementation Plan**: Outline steps for executing the marketing strategy and campaign, including timelines and key milestones.

            7. **Evaluation Metrics**: Specify how to measure the success of the marketing strategy and campaign, focusing on KPIs related to audience engagement, brand awareness, and sales performance.
            
            DON'T USE ANY TOOL."""
            
            
            agent_2 = create_openai_tools_agent(main_model,main_tools, prompt)
            agent_executor_2 = AgentExecutor(agent=agent, tools=main_tools, verbose=True, max_iterations = 500, max_tokens = 3000)
            output = ""
            with st.spinner('Generating marketing strategy and campaign...'):
                output = agent_executor_2.invoke({"input": questionaire_report})
            
            st.write(matter_pyramid + "\n\n" + output['output'])
        else:
            if not uploaded_file:
                st.error(f"Please upload one docx document.")
            if not BRAND_NAME:
                st.error(f"Please enter the brand name.")
            if not INDUSTRY:
                st.error(f"Please enter the industry of your brand.")
